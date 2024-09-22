from docx import Document

def copy_table_preserving_formatting(input_file, output_file):
    # Load the input document
    doc = Document(input_file)
    # Create a new Document for the output
    new_doc = Document()

    # Iterate through each table in the document
    for table in doc.tables:
        # Add a new table to the new document with the same dimensions
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))

        # Copy the content of each cell
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)

                # Create new paragraphs based on the original cell's content
                for paragraph in cell.paragraphs:
                    new_paragraph = new_cell.add_paragraph()
                    # Attempt to preserve the original style
                    try:
                        new_paragraph.style = paragraph.style.name
                    except KeyError:
                        # Fallback to default style if the style is not found
                        new_paragraph.style = new_doc.styles['Normal']
                    
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)
                        # Preserve formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.size = run.font.size
                        new_run.font.color.rgb = run.font.color.rgb

    # Save the new document
    new_doc.save(output_file)

# Usage
input_file = 'inputAP.docx'
output_file = 'outputAP.docx'
copy_table_preserving_formatting(input_file, output_file)
