import pypandoc
from docx import Document
import os
from bs4 import BeautifulSoup

def convert_docx_to_html(input_docx):
    output_html = 'temp_output.html'
    pypandoc.convert_file(input_docx, 'html', outputfile=output_html)
    return output_html

def split_cell_content(cell_text):
    # Split cell content by `|` and move the content to next columns
    parts = cell_text.split('|')
    return parts

def extract_images_and_text(cell, image_files):
    html_content = str(cell)
    if '<img' in html_content:
        return html_content
    else:
        return cell.get_text()

def adjust_table_cells(html_file, image_files):
    with open(html_file, 'r') as file:
        html_content = file.read()

    soup = BeautifulSoup(html_content, 'html.parser')
    tables = soup.find_all('table')
    
    table_adjusted = []
    for table in tables:
        rows = table.find_all('tr')
        for row in rows:
            cells = row.find_all(['td', 'th'])
            adjusted_row = []
            for cell in cells:
                cell_html = extract_images_and_text(cell, image_files)
                if '|' in cell_html:
                    cell_parts = split_cell_content(cell_html)
                    adjusted_row.extend(cell_parts)
                else:
                    adjusted_row.append(cell_html)
            table_adjusted.append(adjusted_row)
    
    return table_adjusted

def create_docx_from_adjusted_table(table_html, output_docx):
    doc = Document()
    
    if not table_html:
        print("No table HTML found.")
        return
    
    # Create a table in the docx file
    num_rows = len(table_html)
    num_cols = max(len(row) for row in table_html)
    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    
    for i, row_html in enumerate(table_html):
        for j, cell_content in enumerate(row_html):
            cell = tbl.cell(i, j)
            if cell_content.startswith('<'):
                # Insert HTML content (i.e., images)
                cell.paragraphs[0].add_run(cell_content)
            else:
                cell.text = cell_content
    
    doc.save(output_docx)

def main(input_docx, output_docx):
    html_file = convert_docx_to_html(input_docx)
    image_files = []  # Update this if you need to manage image files
    table_adjusted = adjust_table_cells(html_file, image_files)
    create_docx_from_adjusted_table(table_adjusted, output_docx)

    # Clean up temporary HTML file
    os.remove(html_file)

if __name__ == "__main__":
    input_docx = 'input.docx'
    output_docx = 'output.docx'
    main(input_docx, output_docx)
