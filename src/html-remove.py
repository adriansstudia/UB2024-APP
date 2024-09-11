from docx import Document
import re

def remove_html_tags(text):
    """Remove HTML tags from text."""
    clean = re.compile('<.*?>')
    return re.sub(clean, '', text)

def process_docx(input_docx, output_docx):
    doc = Document(input_docx)
    new_doc = Document()

    for table in doc.tables:
        # Create a new table with the same number of rows and columns
        num_cols = len(table.rows[0].cells) if table.rows else 0
        new_table = new_doc.add_table(rows=0, cols=num_cols)

        for row in table.rows:
            new_row = new_table.add_row()
            for i, cell in enumerate(row.cells):
                if i >= len(new_row.cells):
                    # Ensure the new row has enough cells
                    new_row.add_cell()
                # Check if cell contains HTML (likely with images)
                if '<img' in cell.text:
                    # Copy the cell content as is if it contains images
                    new_row.cells[i].text = cell.text
                else:
                    # Remove HTML tags if no images are present
                    clean_text = remove_html_tags(cell.text)
                    new_row.cells[i].text = clean_text

    new_doc.save(output_docx)

if __name__ == "__main__":
    input_docx = 'output.docx'
    output_docx = 'cleaned_output.docx'
    process_docx(input_docx, output_docx)
