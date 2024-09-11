from docx import Document

def rearrange_table_columns(input_docx, output_docx, new_order):
    doc = Document(input_docx)
    new_doc = Document()

    for table in doc.tables:
        # Create a new table with the same number of rows and columns
        num_rows = len(table.rows)
        num_cols = len(table.rows[0].cells)
        new_table = new_doc.add_table(rows=num_rows, cols=num_cols)

        # Create a mapping for the current column indices
        header_cells = [cell.text for cell in table.rows[0].cells]
        col_mapping = {header: index for index, header in enumerate(header_cells)}

        # Create a new header row based on new_order
        header_row = new_table.rows[0]
        for i, header in enumerate(new_order):
            header_row.cells[i].text = header

        # Copy data into the new table with the new column order
        for row_index, row in enumerate(table.rows):
            if row_index == 0:
                continue  # Skip the header row, already handled
            new_row = new_table.rows[row_index]
            for col_index, header in enumerate(new_order):
                old_col_index = col_mapping.get(header, None)
                if old_col_index is not None:
                    new_row.cells[col_index].text = row.cells[old_col_index].text
                else:
                    new_row.cells[col_index].text = ''

    new_doc.save(output_docx)

if __name__ == "__main__":
    input_docx = 'cleaned_output.docx'
    output_docx = 'rearranged_output.docx'
    new_order = ["number", "question", "kategoria", "zestaw", "rating", "answer"]
    rearrange_table_columns(input_docx, output_docx, new_order)
