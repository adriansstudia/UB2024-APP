import pypandoc
from docx import Document
import os
import re
import csv
from bs4 import BeautifulSoup

def extract_images(input_docx, image_dir='images'):
    doc = Document(input_docx)
    if not os.path.exists(image_dir):
        os.makedirs(image_dir)

    image_files = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_ext = rel.target_part.content_type.split('/')[1]
            image_name = f'image_{len(image_files) + 1}.{image_ext}'
            image_path = os.path.join(image_dir, image_name)
            
            with open(image_path, 'wb') as img_file:
                img_file.write(image_data)
            
            # Save the full path with the new URL directory
            image_url = f'https://adriansstudia.github.io/UB2024-APP/{image_dir}/{image_name}'
            image_files.append((rel.target_ref, image_url))
    
    return image_files

def convert_docx_to_html(input_docx):
    output_html = 'temp_output.html'
    pypandoc.convert_file(input_docx, 'html', outputfile=output_html, extra_args=['--from=docx', '--to=html'])
    return output_html

def split_cell_content(cell_text):
    # Split cell content by `|` and move the content to next columns
    parts = cell_text.split('|')
    return parts

def extract_images_and_text(cell, image_files):
    html_content = str(cell)
    for ref, url in image_files:
        if ref in html_content:
            # Replace the local path with the URL path
            html_content = html_content.replace(ref, url)
    if '<img' in html_content:
        return html_content
    else:
        return cell.get_text()

def adjust_table_cells(html_file, image_files):
    with open(html_file, 'r', encoding='utf-8') as file:
        html_content = file.read()

    soup = BeautifulSoup(html_content, 'html.parser')
    tables = soup.find_all('table')
    
    table_adjusted = []
    for table in tables:
        rows = table.find_all('tr')
        for row in rows:
            cells = row.find_all(['td', 'th'])
            adjusted_row = []
            for cell in cells:
                cell_html = extract_images_and_text(cell, image_files)
                if '|' in cell_html:
                    cell_parts = split_cell_content(cell_html)
                    adjusted_row.extend(cell_parts)
                else:
                    adjusted_row.append(cell_html)
            table_adjusted.append(adjusted_row)
    
    return table_adjusted

def remove_html_tags(text):
    """Remove HTML tags from text."""
    clean = re.compile('<.*?>')
    return re.sub(clean, '', text)

def create_docx_from_adjusted_table(table_html, output_docx):
    doc = Document()
    
    if not table_html:
        print("No table HTML found.")
        return
    
    # Create a table in the docx file
    num_rows = len(table_html) + 1  # +1 for the heading row
    num_cols = max(len(row) for row in table_html) if table_html else 6  # Assuming 6 columns if no table found
    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    
    # Add heading row
    headings = ["number", "kategoria", "zestaw", "question", "answer", "rating"]
    heading_row = tbl.rows[0]
    for i, heading in enumerate(headings):
        heading_row.cells[i].text = heading
    
    # Add the rest of the rows
    for i, row_html in enumerate(table_html):
        for j, cell_content in enumerate(row_html):
            cell = tbl.cell(i + 1, j)  # Shift rows by 1 to account for heading row
            if cell_content.startswith('<'):
                # Insert HTML content (i.e., images)
                cell.paragraphs[0].add_run(cell_content)
            else:
                cell.text = cell_content
    
    doc.save(output_docx)

def rearrange_table_columns(input_docx, output_docx, new_order):
    doc = Document(input_docx)
    new_doc = Document()

    for table in doc.tables:
        # Create a new table with the same number of rows and columns
        num_rows = len(table.rows)
        num_cols = len(table.rows[0].cells)
        new_table = new_doc.add_table(rows=num_rows, cols=num_cols)

        # Create a mapping for the current column indices
        header_cells = [cell.text for cell in table.rows[0].cells]
        col_mapping = {header: index for index, header in enumerate(header_cells)}

        # Create a new header row based on new_order
        header_row = new_table.rows[0]
        for i, header in enumerate(new_order):
            header_row.cells[i].text = header

        # Copy data into the new table with the new column order
        for row_index, row in enumerate(table.rows):
            if row_index == 0:
                continue  # Skip the header row, already handled
            new_row = new_table.rows[row_index]
            for col_index, header in enumerate(new_order):
                old_col_index = col_mapping.get(header, None)
                if old_col_index is not None:
                    new_row.cells[col_index].text = row.cells[old_col_index].text
                else:
                    new_row.cells[col_index].text = ''

    new_doc.save(output_docx)

def save_table_as_csv(docx_file, csv_file, delimiter=';'):
    doc = Document(docx_file)
    table = doc.tables[0]  # Assuming only one table in the document

    with open(csv_file, 'w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file, delimiter=delimiter)
        
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            writer.writerow(row_data)

def process_docx(input_docx, temp_output_docx, final_output_docx, csv_file):
    image_files = extract_images(input_docx)
    html_file = convert_docx_to_html(input_docx)
    table_adjusted = adjust_table_cells(html_file, image_files)
    create_docx_from_adjusted_table(table_adjusted, temp_output_docx)
    
    # Remove HTML tags from cells that do not contain images
    final_doc = Document(temp_output_docx)
    new_doc = Document()
    
    for table in final_doc.tables:
        # Create a new table with the same number of rows and columns
        num_cols = len(table.rows[0].cells) if table.rows else 0
        new_table = new_doc.add_table(rows=0, cols=num_cols)

        for row in table.rows:
            new_row = new_table.add_row()
            for i, cell in enumerate(row.cells):
                if i >= len(new_row.cells):
                    # Ensure the new row has enough cells
                    new_row.add_cell()
                # Check if cell contains HTML (likely with images)
                if '<img' in cell.text:
                    # Copy the cell content as is if it contains images
                    new_row.cells[i].text = cell.text
                else:
                    # Remove HTML tags if no images are present
                    clean_text = remove_html_tags(cell.text)
                    new_row.cells[i].text = clean_text

    new_doc.save(temp_output_docx)
    
    # Rearrange columns in the final output
    new_order = ["number", "question", "kategoria", "zestaw", "rating", "answer"]
    rearrange_table_columns(temp_output_docx, final_output_docx, new_order)
    
    # Save the rearranged table as CSV
    save_table_as_csv(final_output_docx, csv_file)
    
    # Clean up temporary HTML file
    os.remove(html_file)

if __name__ == "__main__":
    input_docx = 'input.docx'
    temp_output_docx = 'temp_output.docx'
    final_output_docx = 'rearranged_output.docx'
    csv_file = 'output.csv'
    process_docx(input_docx, temp_output_docx, final_output_docx, csv_file)
