from docx import Document
import html  # For escaping special HTML characters

# Define file paths
input_file = "inputW.docx"
output_file = "outputW.docx"

# Function to convert plain text into HTML format
def convert_to_html(paragraphs):
    html_text = []
    for para in paragraphs:
        # Escape special HTML characters and wrap each paragraph in <p> tags
        escaped_text = html.escape(para)
        html_text.append(f"<p>{escaped_text}</p>")
    return "".join(html_text)

# Function to extract headings and normal text
def extract_headings_and_text(doc):
    data = []
    current_question, current_kategoria = "", ""
    normal_text = []
    
    for paragraph in doc.paragraphs:
        style = paragraph.style.name
        text = paragraph.text.strip()

        # Check for heading styles
        if style == 'Heading 2':
            # Save the previous data before resetting, repeat "kategoria" content
            if normal_text or current_kategoria:
                data.append([current_question, current_kategoria, convert_to_html(normal_text)])
            current_question = text
            normal_text = []
        elif style == 'Heading 3':
            # "Kategoria" will repeat till the next heading 3
            current_kategoria = text
        elif text:  # If it's normal text, append it
            normal_text.append(text)
    
    # Append the last set of data
    if normal_text or current_kategoria:
        data.append([current_question, current_kategoria, convert_to_html(normal_text)])
    
    return data

# Function to create output file with table
def create_table_in_docx(data, output_file):
    doc = Document()
    
    # Add a table with 9 columns: number / numberP / question / kategoria / zestaw / rating / answer / aiAnswer / law
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    
    # Add the headers with all lowercase column names
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'number'
    hdr_cells[1].text = 'numberP'
    hdr_cells[2].text = 'question'
    hdr_cells[3].text = 'kategoria'
    hdr_cells[4].text = 'zestaw'
    hdr_cells[5].text = 'rating'
    hdr_cells[6].text = 'answer'
    hdr_cells[7].text = 'aiAnswer'
    hdr_cells[8].text = 'law'
    
    # Add rows to the table
    for row_data in data:
        row_cells = table.add_row().cells
        row_cells[0].text = ""  # number (empty for now)
        row_cells[1].text = ""  # numberP (empty for now)
        row_cells[2].text = row_data[0] if row_data[0] else ""  # question
        row_cells[3].text = row_data[1] if row_data[1] else ""  # kategoria (Heading 3, repeated)
        row_cells[4].text = ""  # zestaw (empty for now)
        row_cells[5].text = ""  # rating (empty for now)
        row_cells[6].text = row_data[2] if row_data[2] else ""  # answer (HTML formatted)
        row_cells[7].text = ""  # aiAnswer (empty for now)
        row_cells[8].text = ""  # law (empty for now)
    
    # Save the document
    doc.save(output_file)

# Main script logic
def main():
    # Load the input document
    doc = Document(input_file)
    
    # Extract the headings and normal text
    data = extract_headings_and_text(doc)
    
    # Create the output document with the table
    create_table_in_docx(data, output_file)
    print(f"Output written to {output_file}")

# Run the script
if __name__ == "__main__":
    main()
